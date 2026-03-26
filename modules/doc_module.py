import os
from tkinter import filedialog, messagebox
import customtkinter as ctk
import subprocess
import shutil
from PIL import Image
import pytesseract
from docx import Document
import fitz
from modules.base_module import (
    BaseModule, BG_DARK, BG_CARD, BG_ITEM,
    ACCENT, ACCENT_H, TEXT_PRI, TEXT_SEC, BORDER
)
from core.job_queue import JobStatus
from ui.widgets import DropZone, FileListWidget

TOOLS = [
    ("merge",   "🔗", "Combinar y Convertir"), # El nuevo panel multiuso
    ("replace", "🔍", "Buscar y Reemplazar"),  # La herramienta de edición en lote
    ("extract", "📄", "Extraer Texto"),        # Para sacar texto plano de cualquier archivo
]

DOC_EXTS = [".docx", ".doc", ".odt", ".txt", ".md", ".rtf"]

class DocsModule(BaseModule):
    module_id = "docs"
    def __init__(self, master, app, **kwargs):
        super().__init__(master, app, **kwargs)
        self.rowconfigure(0, weight=1)

        self.columnconfigure(0, weight=0, minsize=210) # 1. Menú lateral (fijo a 210px)
        self.columnconfigure(1, weight=0)              # 2. Hueco vacío (antiguo divisor)
        self.columnconfigure(2, weight=1)              # 3. PANEL CENTRAL (se estira a tope)

        self._panels = {}
        self._tool_btns = {}
        self._active_tool = ""
        self._build_sidebar()
        self._build_panels()
        self._switch_tool("merge")

    def _build_sidebar(self):
        sb = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0)
        sb.grid(row=0, column=0, sticky="nsew")

        ctk.CTkLabel(sb, text="Herramientas Docs",
                     font=("Segoe UI", 11, "bold"), 
                     text_color=TEXT_SEC).grid(
                     row=0, column=0, padx=14, pady=(18, 10), sticky="w")
        for i, (tid, icon, name) in enumerate(TOOLS):
            btn = ctk.CTkButton(sb, text=f"  {icon}  {name}", anchor="w",
                                fg_color="transparent", hover_color=BG_ITEM,
                                text_color=TEXT_SEC, font=("Segoe UI", 13),
                                height=40, corner_radius=8,
                                command=lambda t=tid: self._switch_tool(t))
            btn.grid(row=i+1, column=0, padx=8, pady=2, sticky="ew")
            self._tool_btns[tid] = btn

    def _switch_tool(self, tool_id):
        if self._active_tool:
            self._panels[self._active_tool].grid_remove()
            self._tool_btns[self._active_tool].configure(fg_color="transparent", text_color=TEXT_SEC)
        self._panels[tool_id].grid()
        self._active_tool = tool_id
        self._tool_btns[tool_id].configure(fg_color=BG_ITEM, text_color=TEXT_PRI)

    def receive_files(self, paths):
        panel = self._panels.get(self._active_tool)
        if panel and hasattr(panel, "allowed_exts"):
            docs = [p for p in paths if os.path.splitext(p)[1].lower() in panel.allowed_exts]
            if docs:
                panel._file_list.add_files(docs)

    def _build_panels(self):
        """Asocia cada ID de herramienta con su clase de panel correspondiente."""
        builders = {
            "merge":   UniversalMergePanel, 
            "replace": ReplacePanel,        # Panel para búsqueda y sustitución masiva
            "extract": ExtractTextPanel,    # Panel dedicado a la extracción de contenido
        }
        
        for tid, cls in builders.items():
            panel = cls(self, app=self.app)
            panel.grid(row=0, column=2, sticky="nsew")
            panel.grid_remove()
            self._panels[tid] = panel


class _BaseDocPanel(ctk.CTkFrame):
    title = ""
    description = ""
    multi_file = False
    allowed_exts = DOC_EXTS

    def __init__(self, master, app, **kwargs):
        super().__init__(master, fg_color=BG_DARK, corner_radius=0, **kwargs)
        self.app = app
        self.columnconfigure(0, weight=1)
        self._build_common()
        self._build_options()
        self._build_run_btn()

    def _build_common(self):
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=24, pady=(22, 4))
        ctk.CTkLabel(hdr, text=self.title,
                     font=("Segoe UI", 17, "bold"), text_color=TEXT_PRI).pack(anchor="w")
        ctk.CTkLabel(hdr, text=self.description,
                     font=("Segoe UI", 11), text_color=TEXT_SEC).pack(anchor="w", pady=(2,0))
        self._drop = DropZone(self, on_files=self._on_drop,
                              extensions=self.allowed_exts, height=100)
        self._drop.grid(row=1, column=0, sticky="ew", padx=24, pady=(12,8))
        self._file_list = FileListWidget(self)
        self._file_list.grid(row=2, column=0, sticky="nsew", padx=24, pady=(0,8))
        self.rowconfigure(2, weight=1)

    def _on_drop(self, paths):
        if not self.multi_file:
            self._file_list.clear()
        self._file_list.add_files(paths)

    def _build_options(self): pass

    def _build_run_btn(self):
        self._run_btn = ctk.CTkButton(self, text="▶  Ejecutar", height=42,
                                      fg_color=ACCENT, hover_color=ACCENT_H,
                                      font=("Segoe UI", 14, "bold"),
                                      corner_radius=10, command=self._run)
        self._run_btn.grid(row=10, column=0, sticky="ew", padx=24, pady=(8,20))

    def _run(self): raise NotImplementedError

    def _on_done(self, job):
        if job.status == JobStatus.DONE:
            r = job.result
            self.app.after(0, lambda: messagebox.showinfo("Listo", f"Completado:\n\n{r}"))
        elif job.status == JobStatus.ERROR:
            e = job.error
            self.app.after(0, lambda: messagebox.showerror("Error", e))

    def _section(self, text):
        f = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=10)
        f.grid(row=3, column=0, sticky="ew", padx=24, pady=(0,10))
        f.columnconfigure(0, weight=1)
        ctk.CTkLabel(f, text=text, font=("Segoe UI", 11, "bold"),
                     text_color=TEXT_SEC).grid(row=0, column=0, padx=14, pady=(10,6), sticky="w")
        return f

class UniversalMergePanel(_BaseDocPanel):
    title = "Combinar y Convertir"
    description = "Mezcla PDFs, DOCX, TXTs y MDs en un único archivo de salida."
    multi_file = True
    allowed_exts = [".docx", ".pdf", ".txt", ".md"]

    def _build_options(self):
        f = self._section("Ajustes de salida")
        opts = ctk.CTkFrame(f, fg_color="transparent")
        opts.grid(row=1, column=0, padx=14, pady=(5, 15), sticky="w")

        ctk.CTkLabel(opts, text="Formato final:", font=("Segoe UI", 12)).grid(row=0, column=0, sticky="w")
        self._target_fmt = ctk.CTkSegmentedButton(opts, values=["PDF", "DOCX", "MD", "TXT"])
        self._target_fmt.set("PDF")
        self._target_fmt.grid(row=0, column=1, padx=15)

    def _run(self):
        paths = self._file_list.paths
        if not paths: return
        
        fmt = self._target_fmt.get().lower()
        out = filedialog.asksaveasfilename(defaultextension=f".{fmt}", 
                                         filetypes=[(fmt.upper(), f"*.{fmt}")])
        if not out: return

        # Aquí llamaríamos a una función '_merge_all_task' que use la lógica universal
        self.app.job_queue.submit(f"Combinando {len(paths)} archivos",
                                 _merge_all_task, paths, out, fmt,
                                 on_done=self._on_done)


class ReplacePanel(_BaseDocPanel):
    title = "Buscar y Reemplazar"
    description = "Modifica el contenido de varios archivos simultáneamente."
    multi_file = True
    allowed_exts = [".docx", ".txt", ".md", ".json", ".xml", ".config", ".ini", ".csv", ".py", ".c", ".cpp"]

    def _build_options(self):
        f = self._section("Parámetros de sustitución")
        
        # Contenedor para alinear los inputs
        opts = ctk.CTkFrame(f, fg_color="transparent")
        opts.grid(row=1, column=0, padx=14, pady=(5, 15), sticky="w")

        # Campo: Buscar
        ctk.CTkLabel(opts, text="Buscar:", font=("Segoe UI", 12)).grid(row=0, column=0, pady=5, sticky="w")
        self._search_ent = ctk.CTkEntry(opts, placeholder_text="Texto original...", width=250)
        self._search_ent.grid(row=0, column=1, padx=(15, 0), pady=5, sticky="w")

        # Campo: Reemplazar
        ctk.CTkLabel(opts, text="Reemplazar con:", font=("Segoe UI", 12)).grid(row=1, column=0, pady=5, sticky="w")
        self._replace_ent = ctk.CTkEntry(opts, placeholder_text="Nuevo texto...", width=250)
        self._replace_ent.grid(row=1, column=1, padx=(15, 0), pady=5, sticky="w")

    def _run(self):
        paths = self._file_list.paths
        if not paths:
            messagebox.showwarning("Sin archivos", "Añade documentos .docx.")
            return

        search = self._search_ent.get()
        replace = self._replace_ent.get()
        
        if not search:
            messagebox.showwarning("Faltan datos", "Debes indicar qué texto quieres buscar.")
            return

        out_dir = filedialog.askdirectory(title="Seleccionar carpeta de salida")
        if not out_dir:
            return

        self.app.job_queue.submit(
            f"Doc: Reemplazando '{search}'",
            _batch_replace_task, paths, out_dir, search, replace,
            on_done=self._on_done
        )

# --- El Panel de la Interfaz ---
class ExtractTextPanel(_BaseDocPanel):
    title = "Extraer Texto (OCR)"
    description = "Extrae texto de documentos digitales, escaneos o imágenes."
    multi_file = True
    allowed_exts = [".docx", ".pdf", ".txt", ".md", ".png", ".jpg", ".jpeg"]

    def _build_options(self):
        f = self._section("Motor de Reconocimiento")
        info = ("Prism usará OCR automáticamente si el PDF es un escaneo.\n"
                "También puedes soltar fotos directamente (.jpg, .png).")
        ctk.CTkLabel(f, text=info, font=("Segoe UI", 11), 
                     text_color=TEXT_SEC, justify="left").grid(
                     row=1, column=0, padx=14, pady=(0, 14), sticky="w")

    def _run(self):
        paths = self._file_list.paths
        if not paths: return
        
        out = filedialog.asksaveasfilename(
            title="Guardar texto extraído en...",
            defaultextension=".txt",
            filetypes=[("Archivo de texto", "*.txt")],
            initialfile="texto_extraido_prism.txt"
        )
        if not out: return

        self.app.job_queue.submit(
            f"Extraer/OCR: {len(paths)} archivo(s)",
            self._proccess_ocr_task, paths, out,
            on_done=self._on_done
        )

    def _proccess_ocr_task(self, paths, output_path, progress_cb=None):
        """Tarea para la cola que une la extracción de todos los archivos."""
        final_content = []
        for i, p in enumerate(paths):
            content = _get_text_universal(p)
            final_content.append(f"--- ARCHIVO: {os.path.basename(p)} ---\n{content}")
            if progress_cb: progress_cb((i+1)/len(paths))
            
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(final_content))
            
        return output_path

def _convert_to_pdf_task(input_path: str, output_path: str, progress_cb=None):
    """Convierte DOCX a PDF priorizando motores nativos para mantener el diseño."""
    ext = os.path.splitext(input_path)[1].lower()
    
    try:
        # --- PLAN A: MS Word (Fidelidad 100%) ---
        if ext in [".docx", ".doc"]:
            try:
                from docx2pdf import convert
                convert(input_path, output_path)
                return f"Convertido con MS Word:\n{output_path}"
            except Exception:
                pass

        # --- PLAN B: LibreOffice (Fidelidad 99%) ---
        soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
        if os.path.exists(soffice_path):
            out_dir = os.path.dirname(output_path)
            cmd = [soffice_path, "--headless", "--convert-to", "pdf", "--outdir", out_dir, input_path]
            subprocess.run(cmd, check=True, capture_output=True)
            
            lo_output = os.path.join(out_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
            if lo_output != output_path and os.path.exists(lo_output):
                shutil.move(lo_output, output_path)
            return f"Convertido con LibreOffice:\n{output_path}"

        # --- PLAN C: Motor Básico de Python (Solo texto, sin diseño) ---
        if ext in [".docx", ".txt", ".md"]:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=11)

            if ext == ".docx":
                doc = Document(input_path)
                lines = [p.text for p in doc.paragraphs]
            else:
                with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                    lines = f.readlines()

            for text in lines:
                clean_text = text.encode('latin-1', 'replace').decode('latin-1').strip()
                if clean_text:
                    pdf.multi_cell(0, 6, clean_text)

            pdf.output(output_path)
            return f"Convertido (Motor Básico):\n{output_path}"
        else:
            return "Formato no soportado por el motor básico. Requiere LibreOffice."
        
    except Exception as e:
        return f"Error en la conversión: {str(e)}"

def _batch_replace_task(paths: list, out_dir: str, search_text: str, replace_text: str, progress_cb=None):
    """Busca y reemplaza texto en múltiples archivos (.docx y texto plano)."""
    count = 0
    errores = []
    
    # Definimos qué extensiones se leen como texto plano ultrarrápido
    TEXT_EXTS = [".txt", ".md", ".json", ".xml", ".config", ".ini", ".csv", ".py", ".c", ".cpp"]

    for i, src in enumerate(paths):
        try:
            ext = os.path.splitext(src)[1].lower()
            stem = os.path.splitext(os.path.basename(src))[0]
            
            # --- RUTA A: Archivos DOCX (Word) ---
            if ext == ".docx":
                doc = Document(src)
                modificado = False
                
                # Buscar en Párrafos y Tablas conservando el formato visual
                for p in doc.paragraphs:
                    for run in p.runs:
                        if search_text in run.text:
                            run.text = run.text.replace(search_text, replace_text)
                            modificado = True
                            
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    if search_text in run.text:
                                        run.text = run.text.replace(search_text, replace_text)
                                        modificado = True
                                        
                if modificado:
                    out = os.path.join(out_dir, f"{stem}_modificado.docx")
                    doc.save(out)
                    count += 1
                    
            # --- RUTA B: Archivos de Texto Plano y Código ---
            elif ext in TEXT_EXTS:
                # Usamos UTF-8 para evitar problemas con tildes o caracteres de código
                with open(src, "r", encoding="utf-8", errors="ignore") as f:
                    contenido = f.read()
                    
                if search_text in contenido:
                    nuevo_contenido = contenido.replace(search_text, replace_text)
                    out = os.path.join(out_dir, f"{stem}_modificado{ext}")
                    
                    with open(out, "w", encoding="utf-8") as f:
                        f.write(nuevo_contenido)
                    count += 1
            else:
                errores.append(f"{os.path.basename(src)}: Formato {ext} no soportado para reemplazar.")

        except Exception as e:
            errores.append(f"{os.path.basename(src)}: {str(e)}")
            
        if progress_cb: progress_cb((i + 1) / len(paths))
        
    # --- RESULTADO ---
    if errores:
        return f"Proceso finalizado con errores.\nModificados: {count}\nFallos:\n" + "\n".join(errores)
    return f"Proceso finalizado.\n{count} archivos generados en:\n{out_dir}"
    
def _merge_all_task(paths: list, output_path: str, target_fmt: str, progress_cb=None):
    """Combina múltiples formatos (PDF, DOCX, TXT, MD) en un único archivo de salida."""
    try:
       # --- CASO A: EL RESULTADO ES UN PDF ---
        if target_fmt == "pdf":
            from pypdf import PdfWriter
            writer = PdfWriter()
            temp_files = [] # Creamos una lista para la "basura"
            
            for i, p in enumerate(paths):
                ext = os.path.splitext(p)[1].lower()
                if ext == ".pdf":
                    writer.append(p)
                else:
                    # Convertimos temporalmente a PDF
                    temp_pdf = p + ".temp.pdf"
                    _convert_to_pdf_task(p, temp_pdf)
                    writer.append(temp_pdf)
                    temp_files.append(temp_pdf) # Apuntamos para borrar luego
                if progress_cb: progress_cb((i+1)/len(paths))
            
            # 1º: Escribimos el archivo final en el disco
            with open(output_path, "wb") as f:
                writer.write(f)
                
            # 2º: AHORA borramos los temporales de forma segura
            for tmp in temp_files:
                if os.path.exists(tmp): os.remove(tmp)

        # --- CASO B: EL RESULTADO ES DOCX / MD / TXT ---
        else:
            merged_text = []
            for i, p in enumerate(paths):
                ext = os.path.splitext(p)[1].lower()
                text = ""
                
                if ext == ".pdf":
                    doc = fitz.open(p)
                    text = "\n".join([page.get_text() for page in doc])
                elif ext == ".docx":
                    doc = Document(p)
                    text = "\n".join([para.text for para in doc.paragraphs])
                else: # TXT / MD
                    with open(p, "r", encoding="utf-8") as f:
                        text = f.read()
                
                merged_text.append(f"--- INICIO ARCHIVO: {os.path.basename(p)} ---\n{text}\n")
                if progress_cb: progress_cb((i+1)/len(paths))
            final_content = "\n\n".join(merged_text)
            if target_fmt == "docx":
                new_doc = Document()
                new_doc.add_paragraph(final_content)
                new_doc.save(output_path)
            else: # TXT / MD
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(final_content)
        return f"Combinado con éxito en {target_fmt.upper()}:\n{output_path}"
    except Exception as e:
        return f"Error combinando: {str(e)}"
    
# Configuramos la ruta del ejecutable de Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def _get_text_universal(path):
    """Extrae texto de archivos digitales o escaneados (OCR)."""
    ext = os.path.splitext(path)[1].lower()
    text = ""

    if ext == ".pdf":

        doc = fitz.open(path)
        # 1. Intento digital (rápido)
        text = "\n".join([page.get_text() for page in doc])
        
        # 2. Si es un escaneo (poco texto digital), activamos OCR
        if len(text.strip()) < 50:
            from pdf2image import convert_from_path
            # Usamos el poppler que ya configuraste antes
            poppler_path = r"C:\poppler\poppler-25.12.0\Library\bin"
            pages = convert_from_path(path, dpi=300, poppler_path=poppler_path)
            
            ocr_parts = []
            for page in pages:
                # 'spa' para español
                ocr_parts.append(pytesseract.image_to_string(page, lang='spa'))
            text = "\n".join(ocr_parts)

    elif ext in [".png", ".jpg", ".jpeg"]:
        # OCR directo para imágenes
        text = pytesseract.image_to_string(Image.open(path), lang='spa')
    
    elif ext == ".docx":
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs])
    
    else: # TXT / MD
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
            
    return text
